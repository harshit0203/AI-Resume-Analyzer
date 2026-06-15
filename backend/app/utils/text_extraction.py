from __future__ import annotations

import io
import re

from docx import Document
from pypdf import PdfReader

from app.core.exceptions import ValidationError

def _normalise(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()

def extract_pdf_text(data: bytes) -> str:
    try:
        reader = PdfReader(io.BytesIO(data))
        pages = [page.extract_text() or "" for page in reader.pages]
    except Exception as exc:
        raise ValidationError("Unable to read the PDF file.") from exc
    return _normalise("\n".join(pages))

def extract_docx_text(data: bytes) -> str:
    try:
        document = Document(io.BytesIO(data))
    except Exception as exc:
        raise ValidationError("Unable to read the DOCX file.") from exc

    parts: list[str] = [para.text for para in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return _normalise("\n".join(parts))

def extract_text(data: bytes, content_type: str, file_name: str) -> str:
    name = file_name.lower()
    if name.endswith(".pdf") or content_type == "application/pdf":
        return extract_pdf_text(data)
    if name.endswith(".docx") or content_type in {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }:
        return extract_docx_text(data)
    raise ValidationError("Unsupported file type. Upload a PDF or DOCX resume.")
